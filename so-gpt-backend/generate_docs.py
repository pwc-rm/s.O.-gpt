"""
Generates 10 synthetic s.Oliver internal documents as DOCX files
and uploads them to Azure Blob Storage (replaces old PwC docs).

Run: python generate_docs.py

Categories:
  HR-Guidelines (4):  Urlaub, Arbeitszeit, Mobiles Arbeiten, Elternzeit/Krank
  IT-Richtlinien (3): Passwortrichtlinie, BYOD, IT-Sicherheit
  Code of Conduct (3): Verhaltenskodex, Geschenke, Datenschutz/Hinweisgeber
"""
from __future__ import annotations
import io
from docx import Document
from docx.shared import Pt, RGBColor
import config
from azure.storage.blob import BlobServiceClient


# ── Helpers ───────────────────────────────────────────────────────────────────

def _doc(title: str, author: str, valid_from: str = "1. Januar 2025") -> Document:
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author
    h = doc.add_heading(title, 0)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph(f"Gültig ab: {valid_from}  ·  Herausgeber: {author}")
    doc.add_paragraph()
    return doc

def h1(doc, text):  doc.add_heading(text, 1)
def h2(doc, text):  doc.add_heading(text, 2)
def h3(doc, text):  doc.add_heading(text, 3)
def body(doc, text): p = doc.add_paragraph(text); p.style.font.size = Pt(10.5)
def bullet(doc, text): doc.add_paragraph(text, style="List Bullet")
def spacer(doc):    doc.add_paragraph()

def table(doc, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers): t.rows[0].cells[i].text = h
    for row in rows:
        r = t.add_row()
        for i, v in enumerate(row): r.cells[i].text = v
    spacer(doc)

def save(doc) -> bytes:
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  HR-GUIDELINES — 4 Dokumente
# ══════════════════════════════════════════════════════════════════════════════

# ── HR 1: Urlaubsregelung ─────────────────────────────────────────────────────
def build_urlaubsregelung() -> bytes:
    doc = _doc("s.Oliver HR-Richtlinie: Urlaubsregelung", "People & Culture, s.Oliver Group GmbH & Co. KG")

    h1(doc, "1. Urlaubsanspruch")
    body(doc, "Alle Mitarbeitenden der s.Oliver Group haben gemäß Tarifvertrag und individuellem Arbeitsvertrag Anspruch auf bezahlten Erholungsurlaub.")
    table(doc, ["Beschäftigungsgruppe", "Urlaubstage/Jahr"], [
        ["Vollzeit Tarif (≥ 3 Jahre Betriebszugehörigkeit)", "30 Arbeitstage"],
        ["Vollzeit Tarif (< 3 Jahre Betriebszugehörigkeit)", "28 Arbeitstage"],
        ["AT-Bereich (Außertariflich)", "30 Arbeitstage"],
        ["Auszubildende (ab 18 Jahre)", "25 Arbeitstage"],
        ["Schwerbehinderte (GdB ≥ 50)", "+5 Zusatztage (gesetzlich)"],
    ])

    h1(doc, "2. Urlaubsantrag und Genehmigung")
    h2(doc, "2.1 Antragstellung über Workday")
    body(doc, "Urlaub wird ausschließlich über das HR-Self-Service-Portal Workday beantragt. Mündliche oder informelle Vereinbarungen sind nicht gültig.")
    table(doc, ["Urlaubsdauer", "Voranmeldefrist"], [
        ["1–3 Tage", "Mindestens 5 Werktage vorher"],
        ["4–9 Tage", "Mindestens 4 Wochen vorher"],
        ["10+ Tage (Urlaub > 2 Wochen)", "Mindestens 8 Wochen vorher"],
        ["Brückentage und Feiertage", "Mindestens 6 Wochen vorher"],
    ])
    h2(doc, "2.2 Genehmigung")
    body(doc, "Die direkte Führungskraft genehmigt den Urlaub. Sie kann die Genehmigung aus betrieblichen Gründen ablehnen (z. B. Jahresabschluss, Messepräsenz, Personalengpass). Ein Rechtsanspruch auf einen bestimmten Urlaubstermin besteht nicht.")

    h1(doc, "3. Sonderurlaub")
    body(doc, "Mitarbeitende haben in folgenden Fällen Anspruch auf bezahlten Sonderurlaub (kein Abzug vom Jahresurlaub):")
    table(doc, ["Anlass", "Tage"], [
        ["Eheschließung (eigene)", "2 Tage"],
        ["Geburt eines Kindes", "2 Tage"],
        ["Tod eines Ehepartners / eingetr. Lebenspartners", "3 Tage"],
        ["Tod eines Elternteils oder Kindes", "2 Tage"],
        ["Tod von Geschwistern / Schwiegereltern", "1 Tag"],
        ["Wohnortwechsel (selber Ort)", "1 Tag"],
        ["Wohnortwechsel (anderer Ort)", "2 Tage"],
        ["Arztbesuche (nur wenn nicht anders möglich)", "Max. 4 Stunden"],
    ])

    h1(doc, "4. Übertrag, Verfall und Abgeltung")
    h2(doc, "4.1 Übertrag")
    body(doc, "Nicht genommener Urlaub aus dem Kalenderjahr kann bis zum 31. März des Folgejahres genommen werden. Danach verfällt er, außer bei:")
    bullet(doc, "nachgewiesener krankheitsbedingter Verhinderung (Urlaubsübertrag bis zu 15 Monaten)")
    bullet(doc, "betrieblich angeordnetem Urlaubsaufschub (Führungskraft bestätigt schriftlich)")
    h2(doc, "4.2 Urlaubsabgeltung bei Ausscheiden")
    body(doc, "Bei Beendigung des Arbeitsverhältnisses werden nicht genommene Urlaubstage in Geld abgegolten, sofern der Urlaub nicht mehr angetreten werden kann. Der Auszahlungsbetrag richtet sich nach dem durchschnittlichen Tagesverdienst der letzten 13 Wochen.")

    h1(doc, "5. Urlaub während Krankheit")
    body(doc, "Erkrankt ein Mitarbeitender während des Urlaubs und liegt eine ärztliche Bescheinigung vor, werden die Krankheitstage nicht als Urlaubstage gewertet. Die Meldung muss unverzüglich beim direkten Vorgesetzten und People & Culture erfolgen.")

    h1(doc, "6. Kontakt")
    body(doc, "Bei Fragen zur Urlaubsregelung: hr@soliver.com oder People & Culture, Uffenheimer Straße 35, 97215 Uffenheim.")
    return save(doc)


# ── HR 2: Arbeitszeit und Überstunden ────────────────────────────────────────
def build_arbeitszeit() -> bytes:
    doc = _doc("s.Oliver HR-Richtlinie: Arbeitszeit und Überstunden", "People & Culture, s.Oliver Group GmbH & Co. KG")

    h1(doc, "1. Regelarbeitszeit")
    table(doc, ["Beschäftigungsgruppe", "Wochenstunden", "Modell"], [
        ["Vollzeit (Tarif)", "37,5 Stunden", "Gleitzeit mit Kernzeit"],
        ["Vollzeit (AT)", "40 Stunden", "Vertrauensarbeitszeit"],
        ["Teilzeit", "Individuell vereinbart", "Gleitzeit oder Festzeit"],
        ["Auszubildende", "37,5 Stunden", "Feste Zeiten je Ausbildungsplan"],
    ])

    h1(doc, "2. Gleitzeitmodell")
    h2(doc, "2.1 Kernarbeitszeiten")
    body(doc, "Alle Mitarbeitenden im Gleitzeitmodell müssen während der Kernarbeitszeiten erreichbar und anwesend sein:")
    table(doc, ["Tag", "Kernarbeitszeit"], [
        ["Montag – Donnerstag", "09:00 – 16:00 Uhr"],
        ["Freitag", "09:00 – 14:00 Uhr"],
    ])
    h2(doc, "2.2 Gleitzeitrahmen")
    body(doc, "Außerhalb der Kernzeiten können Mitarbeitende ihre Arbeitszeit flexibel gestalten:")
    bullet(doc, "Frühester Arbeitsbeginn: 06:30 Uhr")
    bullet(doc, "Spätestes Arbeitsende: 20:00 Uhr")
    bullet(doc, "Mittagspause: Mindestens 30 Minuten (gesetzlich), empfohlen 60 Minuten")
    h2(doc, "2.3 Gleitzeitsaldo")
    body(doc, "Mitarbeitende dürfen ein Gleitzeitsaldo von +40 bis -10 Stunden führen. Bei Überschreitung von +40 Stunden verfallen die Stunden, sofern keine Führungsgenehmigung vorliegt. Negativsaldo > 10 Stunden bedarf der Klärung mit People & Culture.")

    h1(doc, "3. Arbeitszeiterfassung")
    body(doc, "Alle geleisteten Arbeitsstunden sind täglich in Workday zu erfassen. Die Erfassung muss spätestens am Ende der jeweiligen Arbeitswoche abgeschlossen sein.")
    body(doc, "Bei fehlerhafter Erfassung ist die Führungskraft zu informieren. Nachträgliche Korrekturen sind bis zu 4 Wochen rückwirkend möglich.")

    h1(doc, "4. Überstunden")
    h2(doc, "4.1 Definition und Anordnung")
    body(doc, "Überstunden sind Arbeitsstunden, die über die tarifliche oder vertraglich vereinbarte Sollarbeitszeit hinausgehen und von der Führungskraft ausdrücklich angeordnet oder genehmigt wurden. Eigenmächtige Überstunden ohne Genehmigung begründen keinen Anspruch auf Ausgleich.")
    h2(doc, "4.2 Ausgleich")
    table(doc, ["Gruppe", "Ausgleich", "Aufschlag"], [
        ["Tarif-Mitarbeitende", "Freizeitausgleich oder Auszahlung", "25 % Aufschlag bei Auszahlung"],
        ["AT-Mitarbeitende (Ebene 1–4)", "Ausschließlich Freizeitausgleich", "1:1, kein Aufschlag"],
        ["Führungskräfte (Ebene 5+)", "In Grundvergütung enthalten", "Entfällt"],
    ])
    h2(doc, "4.3 Höchstarbeitszeit")
    body(doc, "Gemäß Arbeitszeitgesetz (ArbZG) darf die tägliche Arbeitszeit 10 Stunden nicht überschreiten. Die wöchentliche Arbeitszeit darf 48 Stunden (im 6-Monats-Durchschnitt) nicht übersteigen. Diese Grenzen sind nicht verhandelbar.")

    h1(doc, "5. Nacht-, Sonn- und Feiertagsarbeit")
    body(doc, "Nacht-, Sonn- und Feiertagsarbeit ist nur in begründeten Ausnahmefällen zulässig und bedarf der ausdrücklichen Genehmigung. Zuschläge richten sich nach dem geltenden Tarifvertrag.")
    table(doc, ["Arbeitszeitkategorie", "Zuschlag"], [
        ["Nachtarbeit (23:00–06:00 Uhr)", "25 % oder Freizeitausgleich"],
        ["Sonntagsarbeit", "50 %"],
        ["Feiertagsarbeit", "100 %"],
    ])

    h1(doc, "6. Ruhepausen und Ruhezeiten")
    body(doc, "Gesetzlich vorgeschriebene Mindestpausen (ArbZG § 4):")
    bullet(doc, "Arbeitszeit > 6 Stunden: Pause von mindestens 30 Minuten")
    bullet(doc, "Arbeitszeit > 9 Stunden: Pause von mindestens 45 Minuten")
    body(doc, "Die gesetzliche Mindestruhezeit zwischen zwei Arbeitstagen beträgt 11 Stunden.")
    return save(doc)


# ── HR 3: Mobiles Arbeiten ────────────────────────────────────────────────────
def build_mobiles_arbeiten() -> bytes:
    doc = _doc("s.Oliver HR-Richtlinie: Mobiles Arbeiten (Homeoffice / Remote Work)", "People & Culture, s.Oliver Group GmbH & Co. KG", "1. März 2025")

    h1(doc, "1. Geltungsbereich und Grundsatz")
    body(doc, "Diese Richtlinie gilt für alle Mitarbeitenden der s.Oliver Group in Deutschland, die in einer Position tätig sind, die ortsunabhängiges Arbeiten ermöglicht. Sie basiert auf der Gesamtbetriebsvereinbarung (GBV) Mobiles Arbeiten vom 15. Februar 2024.")
    body(doc, "Mobiles Arbeiten ist ein freiwilliges Angebot von s.Oliver und begründet keinen Rechtsanspruch auf Homeoffice. Die Entscheidung über die Eignung einer Stelle trifft die jeweilige Führungskraft in Absprache mit People & Culture.")

    h1(doc, "2. Anspruch und Umfang")
    h2(doc, "2.1 Maximale Homeoffice-Tage")
    table(doc, ["Beschäftigungsgruppe", "Max. mobile Arbeitstage/Woche", "Mindestpräsenz im Büro"], [
        ["Vollzeit (geeignete Stellen)", "3 Tage", "2 Tage"],
        ["Teilzeit (≥ 20 Std./Woche)", "Anteilig (60 %)", "Anteilig (40 %)"],
        ["Führungskräfte ab Ebene 4", "2 Tage", "3 Tage (Vorbildfunktion)"],
        ["Probezeit (erste 6 Monate)", "1 Tag", "4 Tage"],
        ["Auszubildende", "0 Tage", "5 Tage (Vollpräsenz)"],
    ])
    h2(doc, "2.2 Nicht geeignete Positionen")
    body(doc, "Für folgende Tätigkeiten ist mobiles Arbeiten grundsätzlich nicht möglich:")
    bullet(doc, "Store-Personal, Kassierer, Lager und Logistik")
    bullet(doc, "Empfang, Facility Management, Haustechnik")
    bullet(doc, "Produktionsmitarbeitende")
    bullet(doc, "Tätigkeiten mit Pflicht zur persönlichen Kundenpräsenz")

    h1(doc, "3. Anforderungen an den mobilen Arbeitsplatz")
    h2(doc, "3.1 Technische Voraussetzungen")
    bullet(doc, "Stabile Internetverbindung (min. 25 Mbit/s Download, 5 Mbit/s Upload)")
    bullet(doc, "Nutzung des s.Oliver VPN (GlobalProtect) für alle Arbeitsvorgänge")
    bullet(doc, "Betriebliches Endgerät (Laptop) — private Geräte nur nach BYOD-Genehmigung")
    bullet(doc, "Kamera und Mikrofon für Videokonferenzen (Teams)")
    h2(doc, "3.2 Ergonomie und Arbeitssicherheit")
    body(doc, "Mitarbeitende sind verantwortlich für einen sicheren und ergonomischen Heimarbeitsplatz gemäß Arbeitsstättenverordnung (ArbStättV):")
    bullet(doc, "Ausreichende Beleuchtung (min. 500 Lux am Schreibtisch)")
    bullet(doc, "Ergonomischer Stuhl und Schreibtisch in angemessener Höhe")
    bullet(doc, "Bildschirm auf Augenhöhe, kein Blendlicht")
    bullet(doc, "Abschließbarkeit bei Verarbeitung vertraulicher Daten")
    body(doc, "s.Oliver übernimmt keine Kosten für die Heimarbeitsplatzausstattung (Ausnahme: betrieblich angeordnetes dauerhaftes Homeoffice — separater Prozess über People & Culture).")

    h1(doc, "4. Arbeitsort")
    h2(doc, "4.1 Deutschland")
    body(doc, "Mobiles Arbeiten ist grundsätzlich aus jedem Ort in Deutschland möglich. Der primäre Wohnsitz gilt als Hauptarbeitsort im Homeoffice.")
    h2(doc, "4.2 Europäisches Ausland (Remote Work Abroad)")
    body(doc, "Mobiles Arbeiten aus EU/EWR-Ländern ist bis zu 25 Tage pro Kalenderjahr möglich, sofern:")
    bullet(doc, "die Tätigkeit in das Ausland verlagert wird (kein Urlaub, kein Sabbatical)")
    bullet(doc, "eine Voranmeldung über das Travel & Mobility-Portal erfolgt (min. 4 Wochen vorher)")
    bullet(doc, "People & Culture die steuerliche und sozialversicherungsrechtliche Prüfung abgeschlossen hat")
    bullet(doc, "keine Kundenmeetings vor Ort erforderlich sind")
    h2(doc, "4.3 Nicht-EU-Ausland")
    body(doc, "Arbeit aus Nicht-EU-Ländern ist grundsätzlich nicht gestattet. Ausnahmen (z. B. kurze Dienstreisen mit Homeoffice-Anteil) sind nur mit ausdrücklicher Genehmigung von People & Culture und dem Steuerberater möglich.")

    h1(doc, "5. Erreichbarkeit und Kommunikation")
    body(doc, "Auch im Homeoffice gelten die betrieblichen Kernarbeitszeiten (09:00–16:00 Uhr Mo–Do, 09:00–14:00 Uhr Fr). Mitarbeitende müssen in dieser Zeit:")
    bullet(doc, "telefonisch und per Teams erreichbar sein")
    bullet(doc, "E-Mails innerhalb von 4 Stunden beantworten")
    bullet(doc, "an Teambesprechungen und Videokonferenzen teilnehmen")
    body(doc, "Außerhalb der Kernzeiten haben Mitarbeitende das Recht auf Unerreichbarkeit (Right to disconnect). Führungskräfte sind verpflichtet, dieses Recht zu respektieren.")

    h1(doc, "6. Homeoffice-Vereinbarung")
    body(doc, "Vor Aufnahme des mobilen Arbeitens ist eine individuelle Homeoffice-Vereinbarung über Workday abzuschließen. Die Vereinbarung enthält: Umfang der mobilen Arbeitstage, Arbeitsort, technische Ausstattung, Kündigungsfristen der Vereinbarung. People & Culture stellt das Formular bereit.")
    return save(doc)


# ── HR 4: Elternzeit & Krankmeldung ──────────────────────────────────────────
def build_elternzeit_krank() -> bytes:
    doc = _doc("s.Oliver HR-Richtlinie: Elternzeit, Mutterschutz und Krankmeldung", "People & Culture, s.Oliver Group GmbH & Co. KG")

    h1(doc, "Teil A — Mutterschutz")
    h2(doc, "A.1 Schutzfristen")
    body(doc, "Der gesetzliche Mutterschutz schützt Schwangere und Mütter vor und nach der Geburt:")
    table(doc, ["Phase", "Dauer", "Regelung"], [
        ["Vor der Geburt", "6 Wochen", "Beschäftigungsverbot (außer auf ausdrückl. Wunsch)"],
        ["Nach der Geburt (Regelfall)", "8 Wochen", "Absolutes Beschäftigungsverbot"],
        ["Nach Früh-/Mehrlingsgeburt", "12 Wochen", "Absolutes Beschäftigungsverbot"],
        ["Bei Frühgeburt (< 37. SSW)", "+Differenz zu 8 Wochen", "Verlängertes Beschäftigungsverbot"],
    ])
    h2(doc, "A.2 Mutterschutzlohn")
    body(doc, "Während des Mutterschutzes zahlt s.Oliver Mutterschutzlohn in Höhe des durchschnittlichen Nettoentgelts der letzten 13 Wochen. Die Krankenkasse zahlt einen Festbetrag (€ 13/Tag), s.Oliver trägt den Differenzbetrag zum Nettolohn.")
    h2(doc, "A.3 Schwangerschaftsmitteilung")
    body(doc, "Schwangere Mitarbeitende sind gebeten, die Schwangerschaft frühzeitig People & Culture mitzuteilen — spätestens beim Einsetzen des Mutterschutzes. Die Mitteilung dient ausschließlich dem Schutz der Mitarbeiterin und hat keine negativen Konsequenzen.")

    h1(doc, "Teil B — Elternzeit")
    h2(doc, "B.1 Anspruch und Dauer")
    body(doc, "Eltern haben Anspruch auf Elternzeit bis zur Vollendung des 3. Lebensjahres des Kindes. Beide Elternteile können gleichzeitig Elternzeit nehmen. Ein Anteil von bis zu 24 Monaten kann bis zum 8. Geburtstag des Kindes genommen werden.")
    h2(doc, "B.2 Anmeldung")
    body(doc, "Die Anmeldung muss schriftlich erfolgen — mindestens:")
    bullet(doc, "7 Wochen vor Beginn, wenn Elternzeit innerhalb der ersten 3 Lebensjahre beginnt")
    bullet(doc, "13 Wochen vor Beginn, wenn Elternzeit ab dem 3. Geburtstag genommen wird")
    body(doc, "Die Anmeldung erfolgt über People & Culture (elternzeit@soliver.com). Gleichzeitig ist der Antrag auf Elterngeld beim Familienkasse zu stellen (nicht durch s.Oliver).")
    h2(doc, "B.3 Teilzeit während Elternzeit")
    body(doc, "Mitarbeitende können während der Elternzeit zwischen 15 und 32 Stunden wöchentlich in Teilzeit arbeiten (§ 15 BEEG). Der Antrag ist spätestens 7 Wochen vorher zu stellen. s.Oliver ist verpflichtet, dem Antrag zuzustimmen, sofern dringende betriebliche Gründe nicht entgegenstehen.")
    h2(doc, "B.4 Rückkehr nach Elternzeit")
    body(doc, "Nach Elternzeit haben Mitarbeitende Anspruch auf einen gleichwertigen Arbeitsplatz (gleiche Hierarchieebene, vergleichbare Aufgaben, gleiches Gehalt). Eine erzwungene Versetzung auf eine niedrigere Position ist unzulässig. s.Oliver bietet:")
    bullet(doc, "Halbjährliche Keep-in-touch-Gespräche während der Elternzeit")
    bullet(doc, "Rückkehr-Workshop für Mitarbeitende nach langer Abwesenheit")
    bullet(doc, "Kinderbetreuungszuschuss: 100 €/Monat pro Kind (bis 3. Lebensjahr)")

    h1(doc, "Teil C — Krankmeldung und Arbeitsunfähigkeit")
    h2(doc, "C.1 Meldepflicht")
    body(doc, "Erkrankte Mitarbeitende müssen ihre Arbeitsunfähigkeit unverzüglich, spätestens bis 09:00 Uhr des ersten Krankheitstages, ihrer direkten Führungskraft telefonisch melden. Eine Meldung per E-Mail, SMS oder Teams ist ausreichend nur wenn telefonisch nicht erreichbar.")
    h2(doc, "C.2 Arbeitsunfähigkeitsbescheinigung")
    body(doc, "Seit 2024 gilt die elektronische AU (eAU): Ärzte übermitteln die Krankmeldung direkt an die Krankenkasse. s.Oliver ruft diese ab. Mitarbeitende müssen:")
    bullet(doc, "trotzdem die Führungskraft persönlich informieren (Anruf/Nachricht)")
    bullet(doc, "bei Privatpatienten: AU-Bescheinigung weiterhin in Papierform vorlegen")
    bullet(doc, "ab dem 1. Krankheitstag: AU-Pflicht besteht (kein 'Arbeitgeber-Karenztag')")
    h2(doc, "C.3 Lohnfortzahlung")
    table(doc, ["Zeitraum", "Zahlung", "Von wem"], [
        ["Woche 1–6 (30 Kalendertage)", "100 % des Bruttogehalts", "s.Oliver"],
        ["Ab Woche 7", "Krankengeld ca. 70 % Brutto", "Gesetzliche Krankenkasse"],
        ["Privatpatienten ab Woche 7", "Krankentagegeld (individuell)", "Private Krankenversicherung"],
    ])
    h2(doc, "C.4 Häufige Kurzerkrankungen")
    body(doc, "Bei mehr als 3 Kurzerkrankungen innerhalb von 12 Monaten kann die Führungskraft in Absprache mit People & Culture ein ärztliches Attest ab dem ersten Krankheitstag anfordern. Bevor Maßnahmen eingeleitet werden, findet ein Rückkehrgespräch (Fehlzeitengespräch) statt.")
    h2(doc, "C.5 Betriebliches Eingliederungsmanagement (BEM)")
    body(doc, "Nach mehr als 6 Wochen Arbeitsunfähigkeit innerhalb von 12 Monaten ist s.Oliver gesetzlich verpflichtet, ein BEM-Gespräch anzubieten. Das BEM soll helfen, die Arbeitsfähigkeit zu erhalten. Die Teilnahme ist für Mitarbeitende freiwillig.")
    return save(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  IT-RICHTLINIEN — 3 Dokumente
# ══════════════════════════════════════════════════════════════════════════════

# ── IT 1: Passwortrichtlinie ──────────────────────────────────────────────────
def build_passwortrichtlinie() -> bytes:
    doc = _doc("s.Oliver IT-Richtlinie: Passwort- und Zugriffsmanagement", "IT Security & Compliance, s.Oliver Group GmbH & Co. KG", "1. März 2025")

    h1(doc, "1. Passwortanforderungen")
    body(doc, "Alle Mitarbeitenden der s.Oliver Group sind verpflichtet, starke Passwörter zu verwenden. Die folgenden Mindestanforderungen gelten für alle s.Oliver-Systeme (Active Directory, Microsoft 365, SAP, interne Anwendungen).")
    table(doc, ["Kriterium", "Anforderung", "Erläuterung"], [
        ["Mindestlänge", "12 Zeichen", "Empfohlen: 16+ Zeichen (Passphrase)"],
        ["Großbuchstaben", "Mind. 1", "A–Z"],
        ["Kleinbuchstaben", "Mind. 1", "a–z"],
        ["Ziffern", "Mind. 1", "0–9"],
        ["Sonderzeichen", "Mind. 1", "!@#$%^&*()-_=+[]{}|;':,.<>?"],
        ["Passworthistorie", "Letzten 12 gesperrt", "Kein Wiederholen alter Passwörter"],
        ["Maximale Gültigkeit", "90 Tage", "Automatischer Ablauf-Reminder 14 Tage vorher"],
        ["Kontosperrung", "5 Fehlversuche", "Automatische Sperrung für 30 Minuten"],
    ])

    h1(doc, "2. Verbotene Passwörter und Praktiken")
    h2(doc, "2.1 Schwache Passwörter")
    body(doc, "Folgende Passwortmuster sind systemseitig gesperrt:")
    bullet(doc, "Wörter aus dem Wörterbuch (auch in Kombination mit Zahlen: 'Sommer2024!')")
    bullet(doc, "Namen von Personen, Haustieren oder s.Oliver-Produkten")
    bullet(doc, "Tastaturmuster (qwerty, 123456, asdfgh)")
    bullet(doc, "Wiederholungen (aaa111, abcabc)")
    bullet(doc, "Bekannte kompromittierte Passwörter (wird gegen HaveIBeenPwned geprüft)")
    h2(doc, "2.2 Verbotene Verhaltensweisen")
    bullet(doc, "Weitergabe von Passwörtern an Kollegen, Vorgesetzte oder IT-Support — ausnahmslos verboten")
    bullet(doc, "Aufschreiben auf Papier, Post-it oder in unverschlüsselten Dateien")
    bullet(doc, "Speichern im Browser ohne Verwendung des unternehmenseigenen Passwort-Managers")
    bullet(doc, "Nutzung desselben Passworts für private und berufliche Konten")

    h1(doc, "3. Passwort-Manager")
    body(doc, "s.Oliver stellt allen Mitarbeitenden 1Password for Business kostenlos zur Verfügung. Die Nutzung wird für alle Konten dringend empfohlen.")
    body(doc, "Einrichtung: IT Service Portal (servicedesk.soliver.com) → Kategorie 'Software' → '1Password einrichten'. Die Einrichtung dauert ca. 10 Minuten. Bei Fragen: it-support@soliver.com")

    h1(doc, "4. Multi-Faktor-Authentifizierung (MFA)")
    h2(doc, "4.1 Pflichtpflichtige Systeme")
    body(doc, "MFA ist verpflichtend für:")
    bullet(doc, "Microsoft 365 (Outlook, Teams, SharePoint, OneDrive)")
    bullet(doc, "VPN-Zugang (GlobalProtect)")
    bullet(doc, "Azure und Cloud-Dienste")
    bullet(doc, "SAP (ab Zugriffsklasse 2)")
    bullet(doc, "Alle administrativen Konten")
    h2(doc, "4.2 Zugelassene MFA-Methoden")
    table(doc, ["Methode", "Empfohlen", "Zulässig"], [
        ["Microsoft Authenticator App (Push)", "Ja (bevorzugt)", "Ja"],
        ["TOTP-App (z. B. 1Password, Authy)", "Ja", "Ja"],
        ["SMS-Code", "Nein", "Nur als Fallback"],
        ["E-Mail-Code", "Nein", "Nicht zulässig"],
        ["Hardware-Token (FIDO2 / YubiKey)", "Für Admins", "Ja"],
    ])
    h2(doc, "4.3 Einrichtung")
    body(doc, "Die MFA-Einrichtung erfolgt beim Onboarding durch IT-Support. Bei Gerätewechsel oder -verlust: sofortige Meldung an IT Security (security@soliver.com), damit der alte Faktor deaktiviert wird.")

    h1(doc, "5. Privilegierte und Administrative Konten")
    body(doc, "Für Konten mit erhöhten Rechten (Server-Admin, Datenbankzugriff, Netzwerkverwaltung) gelten zusätzliche Anforderungen:")
    bullet(doc, "Passwortlänge: Mindestens 20 Zeichen")
    bullet(doc, "Passwortrotation: Alle 30 Tage")
    bullet(doc, "Keine Mehrfachnutzung über Systeme hinweg")
    bullet(doc, "Vollständige Protokollierung aller Zugriffe (Audit Log)")
    bullet(doc, "Nur über Privileged Access Workstation (PAW) nutzbar")

    h1(doc, "6. Zugriff beantragen")
    body(doc, "Systemzugänge werden ausschließlich über das IT Service Portal beantragt:")
    body(doc, "servicedesk.soliver.com → Anfragen → Zugriff & Berechtigungen")
    body(doc, "Erforderliche Angaben: System/Anwendung, Business Justification, Berechtigungsstufe, Zeitraum. Jeder Antrag benötigt die Genehmigung der direkten Führungskraft. Kritische Systeme (ERP, Finanz) erfordern zusätzlich IT-Security-Freigabe.")
    return save(doc)


# ── IT 2: BYOD und Gerätenutzung ──────────────────────────────────────────────
def build_byod() -> bytes:
    doc = _doc("s.Oliver IT-Richtlinie: Gerätenutzung und BYOD (Bring Your Own Device)", "IT Security & Compliance, s.Oliver Group GmbH & Co. KG", "1. März 2025")

    h1(doc, "1. Betriebliche Endgeräte")
    h2(doc, "1.1 Ausgabe und Rückgabe")
    body(doc, "s.Oliver stellt Mitarbeitenden in geeigneten Positionen folgende Endgeräte bereit:")
    table(doc, ["Gerät", "Anspruch", "Konfiguration"], [
        ["Laptop (Windows 11)", "Alle Büromitarbeitenden", "Vorinstalliert, MDM-verwaltet"],
        ["Headset (USB-C)", "Mitarbeitende mit Telefonpflicht", "Über IT-Anfrage"],
        ["Smartphone (iOS/Android)", "Führungskräfte ab Ebene 4 + Mobile Roles", "MDM via Intune"],
        ["Monitor (extern)", "Auf Anfrage, Genehmigung FK", "Lieferung über IT-Portal"],
    ])
    body(doc, "Beim Ausscheiden aus dem Unternehmen sind alle betrieblichen Geräte am letzten Arbeitstag an IT-Support zurückzugeben. Nicht zurückgegebene Geräte werden in Rechnung gestellt.")
    h2(doc, "1.2 Wartung und Updates")
    body(doc, "Mitarbeitende sind verpflichtet, Sicherheitsupdates unverzüglich (spätestens binnen 24 Stunden) zu installieren, sobald diese vom IT-System bereitgestellt werden. Geräte mit ausstehenden Updates werden automatisch aus dem Firmennetzwerk ausgesperrt.")

    h1(doc, "2. Private Endgeräte (BYOD)")
    h2(doc, "2.1 Grundsatz")
    body(doc, "s.Oliver erlaubt die eingeschränkte Nutzung privater Smartphones und Tablets für dienstliche Zwecke. Private Laptops dürfen für dienstliche Zwecke nicht genutzt werden (Ausnahme: genehmigter Notfallzugang).")
    h2(doc, "2.2 Zulässige Nutzung auf privaten Geräten")
    table(doc, ["Anwendung", "Erlaubt", "Einschränkung"], [
        ["Microsoft Outlook (mobile App)", "Ja", "Nur s.Oliver E-Mail-Konto"],
        ["Microsoft Teams (mobile App)", "Ja", "Keine vertraulichen Dateien bearbeiten"],
        ["SharePoint (lesend)", "Ja", "Kein Download auf privates Gerät"],
        ["SAP / ERP", "Nein", "Ausnahmslos verboten"],
        ["Kundendaten (CRM)", "Nein", "Ausnahmslos verboten"],
        ["Finanzdaten / Gehaltsdaten", "Nein", "Ausnahmslos verboten"],
    ])
    h2(doc, "2.3 Technische Voraussetzungen für BYOD")
    body(doc, "Vor der Nutzung muss das private Gerät folgende Anforderungen erfüllen:")
    bullet(doc, "Bildschirmsperre: 6-stelliger PIN, Fingerabdruck oder Gesichtserkennung")
    bullet(doc, "Betriebssystem: Nicht älter als 2 Hauptversionen hinter der aktuellen Version")
    bullet(doc, "MDM-App: Microsoft Intune Company Portal muss installiert und registriert sein")
    bullet(doc, "Kein Jailbreak oder Root (führt zur automatischen MDM-Sperrung)")
    bullet(doc, "Automatische Bildschirmsperre nach spätestens 5 Minuten Inaktivität")
    h2(doc, "2.4 Datenschutz auf privaten Geräten")
    body(doc, "s.Oliver erhält durch das MDM (Intune) ausschließlich Zugriff auf den geschäftlichen Container des Geräts. Private Apps, Fotos und Nachrichten sind für s.Oliver nicht einsehbar. Im Verlustfall kann s.Oliver ausschließlich die geschäftlichen Daten remote löschen (Selective Wipe).")
    h2(doc, "2.5 BYOD-Registrierung")
    body(doc, "Registrierung des privaten Geräts: IT Service Portal → BYOD-Antrag → Genehmigung durch Führungskraft → IT-Support sendet Einrichtungslink per E-Mail.")

    h1(doc, "3. Acceptable Use Policy — Nutzungsregeln")
    h2(doc, "3.1 Erlaubte private Nutzung auf Firmengeräten")
    body(doc, "Eine geringfügige private Nutzung von Firmengeräten ist toleriert (z. B. kurzes Nachschauen privater E-Mails in der Mittagspause), sofern:")
    bullet(doc, "die Arbeitsleistung nicht beeinträchtigt wird")
    bullet(doc, "keine Unternehmensressourcen für private Handelsaktivitäten genutzt werden")
    bullet(doc, "keine Daten von Dritten heruntergeladen werden, die Urheberrechte verletzen")
    h2(doc, "3.2 Verbotene Nutzung")
    bullet(doc, "Installation nicht genehmigter Software (nur Software Center / IT-genehmigt)")
    bullet(doc, "Besuch von Glücksspielseiten, illegalen Download-Portalen oder pornografischen Inhalten")
    bullet(doc, "Umgehung von Sicherheitssystemen (Proxy, Firewall, URL-Filter)")
    bullet(doc, "Nutzung von öffentlichen KI-Diensten (ChatGPT, Claude extern) mit vertraulichen Unternehmensdaten")
    bullet(doc, "Mining von Kryptowährungen auf Firmengeräten")
    h2(doc, "3.3 Überwachung")
    body(doc, "s.Oliver behält sich vor, die Nutzung von Firmengeräten und -netzwerken protokollieren zu lassen (Proxy-Logs, E-Mail-Header-Daten, Gerätestatus). Diese Protokollierung dient ausschließlich IT-Sicherheitszwecken und unterliegt den Mitbestimmungsrechten des Betriebsrats.")
    return save(doc)


# ── IT 3: IT-Sicherheit und Datenschutz ──────────────────────────────────────
def build_it_sicherheit() -> bytes:
    doc = _doc("s.Oliver IT-Richtlinie: IT-Sicherheit, Datensicherheit und Vorfallsmeldung", "IT Security & Compliance, s.Oliver Group GmbH & Co. KG", "1. März 2025")

    h1(doc, "1. Datensicherheitsklassen")
    body(doc, "Alle Daten und Dokumente bei s.Oliver werden in folgende Sicherheitsklassen eingeteilt. Die Klassifizierung bestimmt, welche Schutzmaßnahmen anzuwenden sind.")
    table(doc, ["Klassifizierung", "Beispiele", "Erlaubte Speicherorte", "Verschlüsselung"], [
        ["PUBLIC", "Website, Pressemitteilungen, Produktkataloge", "Überall", "Nicht erforderlich"],
        ["INTERNAL", "Interne Memos, Prozesshandbücher, Schulungsunterlagen", "SharePoint, OneDrive", "Optional"],
        ["CONFIDENTIAL", "Gehaltsdaten, Strategiepapiere, Kundendaten, Lieferantenverträge", "SharePoint (gesichert), verschlüsselte Laufwerke", "Pflicht"],
        ["RESTRICTED", "M&A-Dokumente, Quellcode, Vorstandsberichte", "Dedizierte sichere Server, kein Cloud", "Pflicht + Zusatzgenehmigung"],
    ])
    body(doc, "Dokumente ohne Kennzeichnung gelten als INTERNAL. Beim Versand nach außen (E-Mail, USB) muss die Klassifizierung bekannt sein. CONFIDENTIAL und RESTRICTED dürfen nur verschlüsselt übertragen werden.")

    h1(doc, "2. E-Mail-Sicherheit und Phishing")
    h2(doc, "2.1 Phishing erkennen")
    body(doc, "Phishing-E-Mails versuchen, Zugangsdaten zu stehlen oder Schadsoftware zu installieren. Warnsignale:")
    bullet(doc, "Absenderadresse weicht vom erwarteten Domain ab (z. B. @s0liver.com statt @soliver.com)")
    bullet(doc, "Dringlichkeitsappell ('Ihr Konto wird in 24 Stunden gesperrt!')")
    bullet(doc, "Unerwartete Anhänge (ZIP, EXE, Office-Dokumente mit Makros)")
    bullet(doc, "Links, die bei Mouseover auf unbekannte Domains führen")
    bullet(doc, "Schlechte Grammatik oder ungewöhnliche Formatierung")
    h2(doc, "2.2 Umgang mit verdächtigen E-Mails")
    bullet(doc, "Nicht klicken, nicht antworten, keine Anhänge öffnen")
    bullet(doc, "Meldung über Outlook → 'Phishing melden' Button (rechte Maustaste auf E-Mail)")
    bullet(doc, "Alternativ: Weiterleitung an phishing@soliver.com")
    bullet(doc, "Bei versehentlichem Klick: sofort IT Security informieren (security@soliver.com)")

    h1(doc, "3. Umgang mit tragbaren Speichermedien")
    h2(doc, "3.1 USB-Sticks und externe Laufwerke")
    body(doc, "Die Nutzung privater USB-Sticks und externer Laufwerke auf Firmengeräten ist verboten. Dienstlich ausgegebene USB-Sticks (verschlüsselt, Hardware-basiert) sind erlaubt. Anfrage über IT Service Portal.")
    h2(doc, "3.2 CD/DVD, SD-Karten")
    body(doc, "Externe Wechseldatenträger jeglicher Art (außer genehmigten USB-Sticks) sind auf Firmengeräten grundsätzlich gesperrt. Ausnahmen erfordern eine schriftliche Genehmigung von IT Security.")
    h2(doc, "3.3 Cloud-Speicherdienste (privat)")
    body(doc, "Private Cloud-Dienste (Dropbox, Google Drive, iCloud, Wetransfer) dürfen für keine dienstlichen Daten genutzt werden — auch nicht temporär. Für Dateiübertragungen stehen SharePoint und OneDrive (Unternehmensversion) zur Verfügung.")

    h1(doc, "4. Bildschirmsperre und Clean-Desk-Policy")
    h2(doc, "4.1 Bildschirmsperre")
    body(doc, "Der Bildschirm ist zu sperren, wenn der Arbeitsplatz verlassen wird — auch kurzzeitig:")
    bullet(doc, "Windows: Windows-Taste + L")
    bullet(doc, "Mac: Ctrl + Command + Q")
    body(doc, "Firmengeräte sind so konfiguriert, dass sich der Bildschirm nach 5 Minuten Inaktivität automatisch sperrt.")
    h2(doc, "4.2 Clean-Desk-Policy")
    body(doc, "Bei Verlassen des Büros am Ende des Arbeitstages gilt:")
    bullet(doc, "Keine vertraulichen Dokumente offen auf dem Schreibtisch")
    bullet(doc, "Ausdruck von CONFIDENTIAL-Dokumenten: sofort nach Gebrauch schreddern")
    bullet(doc, "Laptop: Herunterfahren oder in verschlossenem Rollcontainer verstauen")
    bullet(doc, "Schlüssel und Zugangskarten: nie sichtbar am Schreibtisch liegen lassen")

    h1(doc, "5. Sicherheitsvorfälle melden")
    h2(doc, "5.1 Was ist ein Sicherheitsvorfall?")
    bullet(doc, "Verlust oder Diebstahl von Endgeräten (Laptop, Handy, USB-Stick)")
    bullet(doc, "Verdacht auf unbefugten Zugriff auf Konten oder Systeme")
    bullet(doc, "Versehentlicher Versand von Daten an falsche Empfänger")
    bullet(doc, "Ransomware-Angriff oder Verschlüsselung von Dateien")
    bullet(doc, "Verdächtige Aktivitäten im Netzwerk oder auf dem Gerät")
    h2(doc, "5.2 Meldekanäle — sofortige Meldung erforderlich")
    table(doc, ["Kanal", "Kontakt", "Verfügbarkeit"], [
        ["IT Security Hotline", "+49 9342 / 883-2580", "24/7 (auch Wochenenden und Feiertage)"],
        ["E-Mail", "security@soliver.com", "Werktags, Antwort innerhalb 2 Stunden"],
        ["IT Service Portal", "servicedesk.soliver.com → Incident", "24/7 Self-Service"],
        ["Datenschutz-Vorfall (DSGVO)", "dsb@soliver.com", "Pflicht bei personenbezogenen Daten"],
    ])
    h2(doc, "5.3 Was tun bei Geräteverlust?")
    body(doc, "1. Sofortige Meldung an IT Security (Hotline), auch außerhalb der Arbeitszeit.")
    body(doc, "2. IT Security sperrt das Gerät remote und löscht ggf. sensible Daten (Remote Wipe).")
    body(doc, "3. Polizeiliche Anzeige bei Diebstahl (Kopie an People & Culture und IT Security).")
    body(doc, "4. Protokollierung des Vorfalls im Security-Portal durch IT Security.")
    return save(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  CODE OF CONDUCT — 3 Dokumente
# ══════════════════════════════════════════════════════════════════════════════

# ── CoC 1: Verhaltenskodex und Interessenkonflikte ───────────────────────────
def build_verhaltenskodex() -> bytes:
    doc = _doc("s.Oliver Code of Conduct: Verhaltenskodex und Interessenkonflikte", "Legal, Compliance & Ethics, s.Oliver Group GmbH & Co. KG")

    h1(doc, "1. Grundwerte der s.Oliver Group")
    body(doc, "Der Code of Conduct der s.Oliver Group ist verbindlich für alle Mitarbeitenden, Führungskräfte und Vorstände weltweit. Er leitet sich aus unseren vier Grundwerten ab:")
    table(doc, ["Wert", "Bedeutung für unser Handeln"], [
        ["Authentizität", "Wir handeln ehrlich und transparent — intern wie extern."],
        ["Zusammenhalt", "Wir respektieren einander und fördern Vielfalt und Inklusion."],
        ["Leidenschaft", "Wir bringen Engagement, Kreativität und Einsatz in unsere Arbeit."],
        ["Verantwortung", "Wir denken nachhaltig und handeln verantwortungsbewusst."],
    ])

    h1(doc, "2. Interessenkonflikte")
    h2(doc, "2.1 Was ist ein Interessenkonflikt?")
    body(doc, "Ein Interessenkonflikt entsteht, wenn private oder persönliche Interessen — oder der Anschein davon — die objektive Wahrnehmung beruflicher Aufgaben bei s.Oliver beeinflussen könnten.")
    h2(doc, "2.2 Typische Situationen")
    bullet(doc, "Nebentätigkeit bei einem Wettbewerber, Lieferanten oder Kunden von s.Oliver")
    bullet(doc, "Beteiligung (> 1 %) an einem Unternehmen, mit dem s.Oliver Geschäftsbeziehungen unterhält")
    bullet(doc, "Entscheidungsbefugnis bei Aufträgen, die an Familienangehörige oder enge Freunde vergeben werden")
    bullet(doc, "Persönliche Annahme von Geschäftsmöglichkeiten, die eigentlich s.Oliver zustehen")
    bullet(doc, "Nutzung von s.Oliver-Informationen (Insiderinformationen) für private Investitionen")
    h2(doc, "2.3 Offenlegungspflicht")
    body(doc, "Mitarbeitende sind verpflichtet, potenzielle Interessenkonflikte unverzüglich schriftlich zu melden — spätestens 5 Werktage nach Kenntnis. Meldung an:")
    bullet(doc, "Direkte Führungskraft (per E-Mail zur Dokumentation)")
    bullet(doc, "Compliance-Abteilung: compliance@soliver.com")
    bullet(doc, "Alternativ: Online-Meldeportal speak-up.soliver.com (auch anonym)")
    body(doc, "Die rechtzeitige Selbstanzeige schützt vor Disziplinarmaßnahmen. Das Compliance-Team entscheidet über Maßnahmen zur Auflösung des Konflikts (z. B. Befangenheitserklärung, Aufgabenteilung, Rotation).")
    h2(doc, "2.4 Nicht offengelegte Konflikte")
    body(doc, "Das bewusste Verschweigen eines Interessenkonflikts stellt eine schwerwiegende Pflichtverletzung dar und kann zur Kündigung führen. Bei materiellem Schaden für s.Oliver kann zusätzlich zivilrechtliche Haftung entstehen.")

    h1(doc, "3. Nebentätigkeiten")
    body(doc, "Die Aufnahme einer Nebentätigkeit (entgeltlich oder unentgeltlich in leitender Funktion) ist grundsätzlich vorab genehmigungspflichtig. Antrag über: hr@soliver.com")
    body(doc, "Genehmigungsvoraussetzungen: keine Beeinträchtigung der Haupttätigkeit, kein Wettbewerb mit s.Oliver, keine Nutzung von s.Oliver-Ressourcen oder -Informationen.")

    h1(doc, "4. Insiderhandelsverbot")
    body(doc, "Mitarbeitende, die Zugang zu kursrelevanten Informationen über s.Oliver oder Geschäftspartner haben, unterliegen dem Insiderhandelsverbot gemäß EU-Marktmissbrauchsverordnung (MAR):")
    bullet(doc, "Kauf oder Verkauf von Wertpapieren auf Basis von Insiderinformationen ist verboten")
    bullet(doc, "Weitergabe von Insiderinformationen an Dritte ist verboten")
    bullet(doc, "Empfehlung von Kauf/Verkauf auf Basis von Insiderinformationen ist verboten")
    body(doc, "Verstöße werden den Strafverfolgungsbehörden gemeldet. Bei Unsicherheit: Compliance-Team kontaktieren bevor gehandelt wird.")

    h1(doc, "5. Wettbewerbsrecht und Kartellrecht")
    body(doc, "s.Oliver hält alle geltenden Wettbewerbs- und Kartellgesetze ein. Verboten sind:")
    bullet(doc, "Absprachen mit Wettbewerbern über Preise, Marktaufteilung oder Mengen")
    bullet(doc, "Missbrauch einer marktbeherrschenden Stellung")
    bullet(doc, "Vereinbarungen mit Lieferanten zur Preisbindung (Resale Price Maintenance)")
    body(doc, "Mitarbeitende in Einkauf, Vertrieb und Strategie müssen jährlich ein Kartellrecht-Training absolvieren (Workday Learning).")
    return save(doc)


# ── CoC 2: Geschenke, Einladungen, Anti-Korruption ───────────────────────────
def build_geschenke_antikorruption() -> bytes:
    doc = _doc("s.Oliver Code of Conduct: Geschenke, Einladungen und Anti-Korruption", "Legal, Compliance & Ethics, s.Oliver Group GmbH & Co. KG")

    h1(doc, "1. Grundsatz Null-Toleranz bei Korruption")
    body(doc, "s.Oliver hat eine absolute Null-Toleranz-Politik gegenüber Korruption und Bestechung in jeder Form. Dies gilt weltweit, unabhängig von lokalen Gepflogenheiten oder Wettbewerbsdruck.")

    h1(doc, "2. Geschenke und Einladungen — Was ist erlaubt?")
    h2(doc, "2.1 Wertgrenzen")
    table(doc, ["Art des Vorteils", "Maximaler Wert", "Bedingungen"], [
        ["Sachgeschenke annehmen", "50 € pro Person und Jahr", "Nicht bei laufenden Vertragsverhandlungen"],
        ["Sachgeschenke geben", "50 € pro Person und Jahr", "Genehmigung FK bei > 30 €"],
        ["Geschäftsessen (Teilnahme)", "100 € pro Person", "Geschäftlicher Anlass, übliche Restaurants"],
        ["Geschäftsessen (Gastgeber)", "150 € pro Person", "FK-Genehmigung, Bewirtungsbeleg erforderlich"],
        ["Veranstaltungstickets (Kultur/Sport)", "150 € pro Ticket", "Genehmigung FK + Compliance-Eintrag"],
        ["Reisen, Übernachtungen (Dritte zahlen)", "Generell verboten", "Ausnahme: Konferenzen (Compliance-Prüfung)"],
    ])
    h2(doc, "2.2 Grenzfälle und Ausnahmen")
    body(doc, "Bei Unklarheiten gilt: Wende dich an compliance@soliver.com, bevor du das Geschenk oder die Einladung annimmst. Eine präventive Anfrage schützt vor Regelverstößen.")

    h1(doc, "3. Geschenkeregister — Meldepflicht")
    body(doc, "Jeder Vorteil, der 30 € übersteigt, ist innerhalb von 5 Werktagen im Geschenkeregister einzutragen:")
    body(doc, "speak-up.soliver.com → Geschenkeregister → Neuen Eintrag erstellen")
    body(doc, "Einzutragende Informationen: Art des Geschenks/Einladung, Schätzwert, Name des Gebers/Empfängers, Datum, Anlass.")
    body(doc, "Nicht gemeldete Vorteile gelten als Regelverstoß und können disziplinarische Folgen haben.")

    h1(doc, "4. Nicht annehmbare Geschenke")
    body(doc, "Folgende Geschenke oder Einladungen dürfen nicht angenommen werden — unabhängig vom Wert:")
    bullet(doc, "Bargeld oder bargeldähnliche Mittel (Gutscheine, Kryptowährungen)")
    bullet(doc, "Geschenke, die bei laufenden Ausschreibungen oder Vertragsverhandlungen angeboten werden")
    bullet(doc, "Geschenke von Anbietern, über deren Aufträge du Entscheidungsbefugnis hast")
    bullet(doc, "Reisen oder Einladungen, die private Anteile enthalten (z. B. Verlängerungstage auf eigene Kosten des Gastgebers)")
    bullet(doc, "Einladungen zu Veranstaltungen sexueller Natur")
    h2(doc, "4.1 Umgang mit unerwarteten Geschenken")
    body(doc, "Wurden Geschenke zugeschickt oder übergeben, bevor abgelehnt werden konnte: Nicht annehmen und höflich zurückgeben, oder beim Compliance-Team melden. Das Team entscheidet über Rückgabe, Weitergabe an Mitarbeitende (Verlosung), Sachspende an gemeinnützige Einrichtung.")

    h1(doc, "5. Bestechung von Amtsträgern")
    body(doc, "Für Zahlungen oder Vorteile gegenüber Amtsträgern (Beamte, Richter, Behördenvertreter) gilt eine noch striktere Null-Toleranz:")
    bullet(doc, "Jede Zahlung an Amtsträger ist verboten — auch wenn sie im betreffenden Land üblich ist")
    bullet(doc, "Sog. 'Facilitation Payments' (kleine Zahlungen zur Beschleunigung von Amtshandlungen) sind ebenfalls verboten")
    bullet(doc, "Selbst geringe Vorteile (< 10 €) an Amtsträger sind zu melden")
    body(doc, "s.Oliver haftet strafrechtlich für Bestechungshandlungen durch Mitarbeitende und Bevollmächtigte weltweit (UK Bribery Act, US FCPA, deutsches StGB § 334).")

    h1(doc, "6. Drittparteien und Agents")
    body(doc, "Die Anti-Korruptionsregeln gelten auch für Dritte, die im Auftrag von s.Oliver handeln:")
    bullet(doc, "Handelsvertreter, Makler, Consultants")
    bullet(doc, "Logistikpartner, Zollagenten, Lobbyisten")
    bullet(doc, "Jeder Dritte, der s.Oliver gegenüber Behörden oder Kunden vertritt")
    body(doc, "s.Oliver führt vor Beauftragung von Drittparteien eine Due-Diligence-Prüfung durch (Compliance-Abteilung koordiniert). Verträge mit Drittparteien enthalten stets Anti-Korruptionsklauseln.")
    return save(doc)


# ── CoC 3: Datenschutz und Hinweisgebersystem ─────────────────────────────────
def build_datenschutz_hinweisgeber() -> bytes:
    doc = _doc("s.Oliver Code of Conduct: Datenschutz, Privatsphäre und Hinweisgebersystem", "Legal, Compliance & Ethics, s.Oliver Group GmbH & Co. KG")

    h1(doc, "1. Datenschutz und DSGVO-Konformität")
    h2(doc, "1.1 Grundsätze")
    body(doc, "s.Oliver verarbeitet personenbezogene Daten von Kunden, Mitarbeitenden, Bewerbern und Geschäftspartnern. Die Einhaltung der DSGVO ist für alle Mitarbeitenden verpflichtend.")
    table(doc, ["DSGVO-Grundsatz", "Bedeutung im Alltag"], [
        ["Rechtmäßigkeit", "Daten nur verarbeiten, wenn Rechtsgrundlage besteht (Vertrag, Einwilligung, berechtigtes Interesse)"],
        ["Zweckbindung", "Daten ausschließlich für den vereinbarten Zweck nutzen — keine Zweckentfremdung"],
        ["Datensparsamkeit", "Nur so viele Daten erheben wie unbedingt nötig"],
        ["Richtigkeit", "Unrichtige Daten unverzüglich berichtigen oder löschen"],
        ["Speicherbegrenzung", "Daten nach Ablauf der Aufbewahrungsfrist löschen"],
        ["Vertraulichkeit", "Unbefugten Zugriff und Datenverlust verhindern"],
    ])
    h2(doc, "1.2 Besonders sensible Daten (Art. 9 DSGVO)")
    body(doc, "Folgende Kategorien bedürfen besonderer Schutzmaßnahmen und dürfen nur mit ausdrücklicher Einwilligung verarbeitet werden:")
    bullet(doc, "Gesundheitsdaten (Krankmeldungen, Behinderungsgrad, AU-Bescheinigungen)")
    bullet(doc, "Biometrische Daten (Fingerabdruck für Zugangssystem)")
    bullet(doc, "Religionszugehörigkeit (relevant für Kirchensteuerabzug)")
    bullet(doc, "Politische Meinungen und Gewerkschaftszugehörigkeit")
    bullet(doc, "Genetische Daten, sexuelle Orientierung")
    h2(doc, "1.3 Aufbewahrungsfristen")
    table(doc, ["Datenart", "Aufbewahrungsfrist", "Rechtsgrundlage"], [
        ["Lohn- und Gehaltsunterlagen", "10 Jahre", "HGB / AO"],
        ["Personalakte (allgemein)", "3 Jahre nach Austritt", "§ 195 BGB"],
        ["Bewerberdaten (abgelehnte Bewerber)", "6 Monate", "AGG"],
        ["Videoüberwachung (Empfang etc.)", "72 Stunden", "Datenschutzbehörde"],
        ["E-Mail-Korrespondenz (dienstlich)", "6 Jahre", "HGB"],
    ])
    h2(doc, "1.4 Datenpannen (Data Breach)")
    body(doc, "Bei Verlust oder unberechtigtem Zugriff auf personenbezogene Daten:")
    bullet(doc, "Innerhalb von 24 Stunden: Meldung an Datenschutzbeauftragten (dsb@soliver.com)")
    bullet(doc, "Innerhalb von 72 Stunden: Meldung an Datenschutzbehörde (sofern Risiko für Betroffene)")
    bullet(doc, "Ggf. Information der Betroffenen (falls hohes Risiko für deren Rechte und Freiheiten)")
    body(doc, "Das Datenschutzteam koordiniert alle Meldungen. Mitarbeitende müssen umgehend informieren — keine Eigeninitiative bei der behördlichen Meldung.")

    h1(doc, "2. Hinweisgebersystem (Whistleblowing)")
    h2(doc, "2.1 Zweck")
    body(doc, "s.Oliver betreibt ein vertrauliches Hinweisgebersystem, das es Mitarbeitenden ermöglicht, Verstöße gegen den Code of Conduct, gesetzliche Vorschriften oder ethische Grundsätze sicher zu melden — ohne Angst vor Konsequenzen.")
    h2(doc, "2.2 Meldbare Verstöße")
    bullet(doc, "Verstöße gegen den Code of Conduct oder interne Richtlinien")
    bullet(doc, "Korruption, Bestechung, Unterschlagung")
    bullet(doc, "Verstöße gegen Datenschutz, Arbeitssicherheit oder Umweltrecht")
    bullet(doc, "Finanzielle Unregelmäßigkeiten oder Bilanzbetrug")
    bullet(doc, "Diskriminierung, Belästigung oder Mobbing")
    bullet(doc, "Verstöße gegen Exportkontroll- oder Sanktionsrecht")
    h2(doc, "2.3 Meldekanäle")
    table(doc, ["Kanal", "Kontakt", "Anonym möglich"], [
        ["Online-Plattform", "speak-up.soliver.com", "Ja (bevorzugt)"],
        ["Hotline (DE)", "0800 / 777 2345 (kostenlos, 24/7)", "Ja"],
        ["E-Mail", "ethics@soliver.com", "Nein"],
        ["Compliance Officer direkt", "Chief Compliance Officer, Uffenheim", "Nein"],
        ["Externe Meldebehörde", "Bundesamt für Justiz (HinSchG)", "Ja"],
    ])
    h2(doc, "2.4 Schutz der Hinweisgebenden")
    body(doc, "s.Oliver garantiert umfassenden Schutz gemäß Hinweisgeberschutzgesetz (HinSchG):")
    bullet(doc, "Keinerlei Vergeltungsmaßnahmen (Kündigung, Versetzung, Gehaltskürzung, Mobbing)")
    bullet(doc, "Wahrung der Vertraulichkeit — Identität wird ohne Zustimmung nicht offengelegt")
    bullet(doc, "Schadensersatzpflicht von s.Oliver bei unzulässigen Vergeltungsmaßnahmen")
    bullet(doc, "Beweislastumkehr: s.Oliver muss nachweisen, dass Maßnahme kein Retaliation war")
    h2(doc, "2.5 Falschmeldungen")
    body(doc, "Mutwillige Falschmeldungen, die einer anderen Person schaden, sind selbst ein Verstoß gegen den Code of Conduct und können disziplinarische Konsequenzen haben. Gutgläubige Meldungen sind jedoch immer geschützt — auch wenn sich ein Verdacht als unbegründet herausstellt.")

    h1(doc, "3. Antidiskriminierung und respektvolles Miteinander")
    h2(doc, "3.1 Geschützte Merkmale (AGG § 1)")
    body(doc, "Diskriminierung aufgrund folgender Merkmale ist verboten:")
    table(doc, ["Merkmal", "Beispielhafte verbotene Handlungen"], [
        ["Rasse und ethnische Herkunft", "Schlechtere Beurteilung, Beleidigungen, Ausgrenzung"],
        ["Geschlecht / Geschlechtsidentität", "Unterschiedliche Bezahlung ohne sachlichen Grund, trans-feindliche Kommentare"],
        ["Religion oder Weltanschauung", "Verweigerung von Gebetszeiten (soweit betrieblich möglich)"],
        ["Behinderung", "Verweigerung angemessener Vorkehrungen, Mobbing"],
        ["Alter", "Stellenausschreibungen mit Altersgrenzen ohne sachlichen Grund"],
        ["Sexuelle Orientierung", "Homophobe Äußerungen, Benachteiligung bei Beförderung"],
    ])
    h2(doc, "3.2 Sexuelle Belästigung")
    body(doc, "Sexuelle Belästigung am Arbeitsplatz (verbale, nonverbale oder körperliche Handlungen sexueller Natur) ist eine schwerwiegende Pflichtverletzung.")
    body(doc, "Betroffene können sich wenden an: hr-confidential@soliver.com (vertraulich), eine Vertrauensperson (betriebliche Vertrauensstelle), den Betriebsrat oder die externe Beschwerdestelle.")
    h2(doc, "3.3 Konsequenzen bei Verstößen")
    table(doc, ["Schweregrad", "Mögliche Konsequenzen"], [
        ["Leicht (erster Vorfall, geringer Schaden)", "Ermahnung, Pflicht-Training, Klärungsgespräch"],
        ["Mittel (Wiederholung oder erheblicher Schaden)", "Schriftliche Abmahnung, Versetzung"],
        ["Schwer (schwerwiegend oder strafrechtlich relevant)", "Fristlose Kündigung, Strafanzeige"],
    ])
    return save(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  UPLOAD
# ══════════════════════════════════════════════════════════════════════════════

DOCUMENTS = [
    # HR-Guidelines
    ("s.Oliver_Urlaubsregelung.docx",                  build_urlaubsregelung,          "HR",  "Richtlinie", "Urlaub"),
    ("s.Oliver_Arbeitszeit-und-Ueberstunden.docx",     build_arbeitszeit,              "HR",  "Richtlinie", "Arbeitszeit"),
    ("s.Oliver_Mobiles-Arbeiten.docx",                 build_mobiles_arbeiten,         "HR",  "Richtlinie", "Mobiles Arbeiten"),
    ("s.Oliver_Elternzeit-und-Krankmeldung.docx",      build_elternzeit_krank,         "HR",  "Richtlinie", "Elternzeit"),
    # IT-Richtlinien
    ("s.Oliver_Passwortrichtlinie.docx",               build_passwortrichtlinie,       "IT",  "Richtlinie", "Passwort"),
    ("s.Oliver_BYOD-und-Geraetenutzung.docx",          build_byod,                     "IT",  "Richtlinie", "BYOD"),
    ("s.Oliver_IT-Sicherheit.docx",                    build_it_sicherheit,            "IT",  "Richtlinie", "IT-Sicherheit"),
    # Code of Conduct
    ("s.Oliver_Verhaltenskodex.docx",                  build_verhaltenskodex,          "CoC", "Richtlinie", "Verhaltenskodex"),
    ("s.Oliver_Geschenke-und-Antikorruption.docx",     build_geschenke_antikorruption, "CoC", "Richtlinie", "Anti-Korruption"),
    ("s.Oliver_Datenschutz-und-Hinweisgebersystem.docx", build_datenschutz_hinweisgeber, "CoC", "Richtlinie", "Datenschutz"),
]

OLD_BLOBS = [
    "GBV Urlaubsgrundsätze.pdf",
    "IT-Richtlinie.pdf",
    "PwC GmbH WPG_GBV Smart Working.pdf",
    "PwC GmbH WPG_GBV Reward.pdf",
    "PwC Legal AG RAG_GBV Bring Your Own Device.pdf",
]

def main():
    bsc = BlobServiceClient.from_connection_string(config.BLOB_CONNECTION_STRING)
    cont = bsc.get_container_client(config.BLOB_CONTAINER_NAME)

    print("─── Deleting old blobs ──────────────────────────────────")
    for name in OLD_BLOBS:
        try:
            cont.delete_blob(name)
            print(f"  ✓ Deleted: {name}")
        except Exception as e:
            print(f"  ⚠ Could not delete '{name}': {e}")

    print("\n─── Uploading 10 s.Oliver documents ────────────────────")
    for filename, builder_fn, *_ in DOCUMENTS:
        data = builder_fn()
        cont.upload_blob(name=filename, data=data, overwrite=True)
        print(f"  ✓ {filename}  ({len(data):,} bytes)")

    print("\nDone. Run ingestion next:\n  python ingestion.py")

if __name__ == "__main__":
    main()
